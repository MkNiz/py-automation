import docx

dx = docx.Document()

dx.add_heading("Check out this cat", 0)
dx.add_picture('kitty.jpg', width=docx.shared.Inches(3), height=docx.shared.Cm(5))
dx.add_heading("Isn't it a good cat?", 1)

dx.save("cat_doc.docx")
print("Document saved.")