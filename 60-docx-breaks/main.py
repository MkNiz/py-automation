import docx

dx = docx.Document()

p1 = dx.add_paragraph("Henlo")
# Line Break
p1.runs[0].add_break()
p2 = dx.add_paragraph("There was a line break up there")
# Page Break
p2.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
p2.add_run("wow that was a page break alright")

dx.save("broken_doc.docx")
print("Document saved")