import docx

dx = docx.Document('a_docx.docx')

print("Paragraphs:", len(dx.paragraphs))
for ix, p in enumerate(dx.paragraphs):
    print("Paragraph " + str(ix+1) + ": " + p.text)
    print("\tRuns: ", len(p.runs))
    for ix, r in enumerate(p.runs):
        print("\t\t" + str(ix+1) + ".) " + r.text)