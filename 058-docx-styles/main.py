import docx

myDoc = docx.Document('a_docx.docx')

# Setting a paragraph style
print("Current style of Paragraph 1:", myDoc.paragraphs[0].style.name)
myDoc.paragraphs[0].style = 'good'
print("New style of Paragraph 1:", myDoc.paragraphs[0].style.name)

# Setting a character style
myDoc.paragraphs[1].runs[0].style = 'excellent'
print("New style of first run of Paragraph 2:", myDoc.paragraphs[1].runs[0].style.name)
print("")

# NOTE: Both paragraph and character styles should exist in the document

myDoc.save("a_restyled_doc.docx")
print("Document saved.")