import docx

dx = docx.Document()

# Add headings
dx.add_heading('My Nice Document', 0)
dx.add_heading('Welcome to my Digital World', 1)

# Add a paragraph
dx.add_paragraph("I'm the first paragraph")

# Add a paragraph and return its object
p2 = dx.add_paragraph("I'm the second paragraph. ")
p2.add_run("I'm a run that was added to the second paragraph.")

dx.save("my_generated_doc.docx")
print("Document saved.")